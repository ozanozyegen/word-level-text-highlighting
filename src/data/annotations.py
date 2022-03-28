from collections import defaultdict
import os, re
import pandas as pd
import random
import gensim

import docx
from docx import *
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from data.helpers import Chat, Paragraph
from features.tokenizer import Tokenizer
from smart_reply.reply_model import SmartReplyModel

class Annotations:
    """ Class for processing doctor-patient conversations """
    def __init__(self, config, PATH=None, first_n=None, tokenizer=Tokenizer(),
        is_tuning_train=False, is_tuning_val=False):
        """Read annotated chats from word files

        Args:
            config (dict): Exp config
            PATH (str): Folder path
            first_n (int, optional): Only take first N, for fast run
            tokenizer (Tokenizer, optional): Custom tokenizer, load new one as default
            is_tuning_train (bool, optional): Use train set
            is_tuning_val (bool, optional): Use validation set
        """
        self.config = config
        self._PATH = PATH
        self._file_ids = os.listdir(self._PATH)[:first_n]
        if is_tuning_train or is_tuning_val: # 80/20 split of the train data
            random.seed(0) # Files will always be shuffled the same way
            random.shuffle(self._file_ids)
            split_idx = int(len(self._file_ids)*0.8)
            if is_tuning_train:
                self._file_ids = self._file_ids[:split_idx]
            elif is_tuning_val:
                self._file_ids = self._file_ids[split_idx:]

        if config.get('debug', False): # Use a subset of the dataset
            self._file_ids = self._file_ids[:10]
        self._tokenizer = tokenizer
        if PATH:
            self._chat_dicts = self.read_chats()
            self.preprocess_files() # Preprocessing and tokenization


    def read_chats(self):
        """Read chats from word files

        Returns:
            chat_dicts (dict): list of dicts where keys are the file ids and
            values are the chat_dict(s)
        """
        chat_dicts = {}
        random.seed(0)
        random.shuffle(self._file_ids)
        for idx, file_id in enumerate(self._file_ids):
            doc_path = os.path.join(self._PATH, file_id)
            chat_dict = self.read_chat(doc_path)
            chat_dicts[file_id] = chat_dict
        return chat_dicts

    def read_chat(self, DOC_PATH):
        """ Read a single chat document
        Args:
            DOC_PATH (str): path for the chat doc
        Returns:
            chat_dict (dict): Chat.content stores chat paragraphs as list of dicts
        """
        document = Document(DOC_PATH)
        chat_id = re.search("([\S]*):\|\|([\s\S]*?)\|\|", document.paragraphs[0].text).group(2)
        chat_dict = {Chat.id: chat_id, Chat.content: []}
        for idx, paragraph in enumerate(document.paragraphs[1:]):
            if idx % 2 == 0:
                person = paragraph.text.strip('\n').strip('||').strip(':')
            else:
                para_dict = {Paragraph.person:person, Paragraph.unproc_tokens:[], Paragraph.unproc_tokens_hi:[]}
                for word in paragraph.runs:
                    splits = word.text.strip('||').split()
                    for split in splits:
                        highlighted = int(word.font.highlight_color == WD_COLOR_INDEX.YELLOW)
                        para_dict[Paragraph.unproc_tokens].append(split)
                        para_dict[Paragraph.unproc_tokens_hi].append(highlighted)
                chat_dict[Chat.content].append(para_dict)
        return chat_dict

    def preprocess_files(self):
        """ Tokenize each paragraph
        """
        for file_id in self._file_ids:
            chat_dict = self._chat_dicts[file_id]
            for para_dict in chat_dict[Chat.content]:
                sentence = " ".join(para_dict[Paragraph.unproc_tokens])
                para_dict[Paragraph.proc_tokens], \
                para_dict[Paragraph.proc_tokens_rem], \
                para_dict[Paragraph.proc_tokens_hi] = self._tokenizer.tokenize(para_dict[Paragraph.unproc_tokens],
                                                                                para_dict[Paragraph.unproc_tokens_hi])

    def get_documents_and_highlights(self,):
        """ Returns documents and highlights as list of lists
        """
        annot_data, annot_data_hi = [], []
        for file_id in self._file_ids:
            for para_dict in self._chat_dicts[file_id][Chat.content]:
                words = para_dict[Paragraph.proc_tokens_rem]
                highlights = para_dict[Paragraph.proc_tokens_hi]

                annot_data.append(" ".join(words))
                annot_data_hi.append(highlights)
        return annot_data, annot_data_hi

    def get_documents(self):
        def get_word_docs(file_ids):
            """ Return chats as dict of lists
            """
            docs_dict = {}
            for file_id in file_ids:
                doc = [para_dict[Paragraph.proc_tokens_rem] for para_dict in self._chat_dicts[file_id][Chat.content]]
                docs_dict[file_id] = doc
            return docs_dict
        return get_word_docs(self._file_ids)

    def report_results(self, chat_dicts,
                       WRITE_DIR='reports/tf-idf/',
                       REPORT_DEBUG=False,):
        """ Write the chat_dicts predicted highlights into a directory as word files
        for detailed analysis of the model predictions

        Args:
            chat_dicts (dict):
            WRITE_DIR (str, optional): Write directory
            REPORT_DEBUG (bool, optional): If true, highlight False Positives and True Negatives in blue and red
        """
        if not os.path.exists(WRITE_DIR):
            os.makedirs(WRITE_DIR)
        for file_id in chat_dicts.keys():
            chat_dict = chat_dicts[file_id]
            document = Document()
            document.add_paragraph("Chat id:||" + chat_dict[Chat.id] + '||')
            for para_dict in chat_dict[Chat.content]:
                document.add_paragraph(f'\n\n||{para_dict[Paragraph.person]}:||')
                paragraph = document.add_paragraph()
                tokens = para_dict[Paragraph.unproc_tokens]
                highlights = self._tokenizer._map_preds(para_dict[Paragraph.proc_tokens],
                    para_dict[Paragraph.pred_hi], para_dict[Paragraph.unproc_tokens])
                trues = para_dict[Paragraph.unproc_tokens_hi]

                for word, hi, true in zip(tokens, highlights, trues):
                    run = paragraph.add_run(word+' ')
                    if hi: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if REPORT_DEBUG:
                        if not hi and true: # True highlight, prediction is not highlight, FN
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                        if hi and not true: # Prediction is highlight, true is not highlight - FP
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE

            document.save(os.path.join(WRITE_DIR, file_id))

    def count_messages_per_chat(self, chat_dicts, WRITE_DIR,
                                file_name='count_messages',
                                only_patient=True):
        """ Count messages per chat into a pandas dataframe
        Args:
            chat_dicts (dict):
            WRITE_DIR (bool):
            file_name (str, optional):
            only_patient (bool, optional): Count only patient messages
        """
        message_counter = defaultdict(lambda: 0)
        for file_id in chat_dicts.keys():
            chat_dict = chat_dicts[file_id]
            for para_dict in chat_dict[Chat.content]:
                if only_patient and para_dict[Paragraph.person] == 'Patient':
                   message_counter[file_id]  += 1

        df = pd.DataFrame(message_counter.items())
        if not os.path.exists(WRITE_DIR):
            os.makedirs(WRITE_DIR)
        df.to_csv(os.path.join(WRITE_DIR, f'{file_name}.csv'))







