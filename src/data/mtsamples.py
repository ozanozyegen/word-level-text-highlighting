""" Code for processing the mtsamples dataset, works similarly
    to annotations.py
"""
import os, re, random
import pandas as pd
import abc
from abc import abstractmethod

import docx
from docx import *
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from data.helpers import Chat, Paragraph
from features.tokenizer import Tokenizer


class Dataset(abc.ABC):

    @abstractmethod
    def _create_dataset(self,):
        pass


class MtsamplesDataset(Dataset):
    def __init__(self, config,
        PATH=None, first_n=None, tokenizer=Tokenizer(), # Params for loading dataset
        create_dataset=False, DATASET_PATH="data/raw/mtsamples.csv"): # Params for creating dataset
        self.config = config
        self._PATH = PATH
        if create_dataset: self._create_dataset(DATASET_PATH)

        self._file_ids = os.listdir(self._PATH)[:first_n]
        self._tokenizer = Tokenizer()
        # Read files similarly to the Annotations class, variables named similarly
        self._chat_dicts = self._read_chats()
        self._preprocess_files()

    def _read_chats(self):
        chat_dicts = {}
        random.seed(0)
        random.shuffle(self._file_ids)
        for file_id in self._file_ids:
            doc_path = os.path.join(self._PATH, file_id)
            chat_dict = self._read_chat(doc_path)
            chat_dicts[file_id] = chat_dict
        return chat_dicts

    def _read_chat(self, DOC_PATH):
        document = Document(DOC_PATH)
        chat_dict = {Chat.content: []}
        for idx, paragraph in enumerate(document.paragraphs):
            if paragraph.text and idx % 2 == 0:
                row_idx = paragraph.text.split('||')[-2]
            else:
                para_dict = {'row_id':row_idx, Paragraph.person: 'Patient', Paragraph.unproc_tokens:[], Paragraph.unproc_tokens_hi:[]}
                for word in paragraph.runs:
                    splits = word.text.strip('||').split()
                    for split in splits:
                        if 'Description||' in split: continue;
                        highlighted = int(word.font.highlight_color == WD_COLOR_INDEX.YELLOW)
                        para_dict[Paragraph.unproc_tokens].append(split)
                        para_dict[Paragraph.unproc_tokens_hi].append(highlighted)
                chat_dict[Chat.content].append(para_dict)
        return chat_dict

    def _preprocess_files(self,):
        for file_id in self._file_ids:
            chat_dict = self._chat_dicts[file_id]
            for para_dict in chat_dict[Chat.content]:
                para_dict[Paragraph.proc_tokens], \
                para_dict[Paragraph.proc_tokens_rem], \
                para_dict[Paragraph.proc_tokens_hi] = self._tokenizer.tokenize(para_dict[Paragraph.unproc_tokens],
                                                                                para_dict[Paragraph.unproc_tokens_hi])

    def _create_dataset(self, DATASET_PATH, rows_per_file=50):
        if not os.path.exists(self._PATH): os.mkdir(self._PATH)
        df = pd.read_csv(DATASET_PATH)
        for start_row_idx in range(0, len(df), rows_per_file):
            self._write_texts_to_file(df, start_row_idx, rows_per_file)

    def _write_texts_to_file(self, df, start_row_idx, rows_per_file):
        document = Document()
        for row_idx in range(start_row_idx, start_row_idx+rows_per_file):
            if row_idx < len(df):
                row = df.iloc[row_idx]
                document.add_paragraph("Row id:||" + str(row_idx) + '||')
                document.add_paragraph("Description||" + row['description'] + '||')
        document.save(os.path.join(self._PATH, f'{row_idx}.docx'))

    def report_results(self, chat_dicts, WRITE_DIR='reports/mtsamples/', REPORT_DEBUG=False):
        """ Write the chat_dicts predicted highlights into a directory
        DEBUG: If true, highlight False Positives and True Negatives in blue and red
        """
        if not os.path.exists(WRITE_DIR):
            os.makedirs(WRITE_DIR)
        for file_id in chat_dicts.keys():
            chat_dict  = chat_dicts[file_id]
            document = Document()
            for para_dict in chat_dict[Chat.content]:
                row_idx = para_dict['row_id']
                document.add_paragraph("Row id:||" + str(row_idx) + '||')

                paragraph = document.add_paragraph()
                tokens = para_dict[Paragraph.unproc_tokens]
                highlights = self._tokenizer._map_preds(para_dict[Paragraph.proc_tokens],
                    para_dict[Paragraph.pred_hi], para_dict[Paragraph.unproc_tokens])
                trues = para_dict[Paragraph.unproc_tokens_hi]

                for word, hi, true in zip(tokens, highlights, trues):
                    run = paragraph.add_run(word+' ')
                    if hi: run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if REPORT_DEBUG:
                        if not hi and true: # True highlight, prediction is not highlight, FN
                            run.font.highlight_color = WD_COLOR_INDEX.RED
                        if hi and not true: # Prediction is highlight, true is not highlight - FP
                            run.font.highlight_color = WD_COLOR_INDEX.BLUE
            document.save(os.path.join(WRITE_DIR, file_id))

    def get_documents(self):
        samples_dict = {}
        for file_id in self._chat_dicts.keys():
            chat_dict  = self._chat_dicts[file_id]
            for para_dict in chat_dict[Chat.content]:
                samples_dict[para_dict['row_id']] = [para_dict[Paragraph.proc_tokens_rem]]

        return samples_dict
